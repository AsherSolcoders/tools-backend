"""PDF tool processors (pypdf for structural ops, PyMuPDF for rendering/encryption)."""
from __future__ import annotations

import io
import mimetypes
from pathlib import Path

from app.core.temp_files import new_result_path
from app.tools.registry import ResultFile, ToolResult, register


def _save_result(data: bytes, name: str) -> ResultFile:
    path = new_result_path(name)
    path.write_bytes(data)
    return ResultFile(token=path.name, name=name, size=len(data), mime="application/pdf")


def _save_named(data: bytes, name: str) -> ResultFile:
    path = new_result_path(name)
    path.write_bytes(data)
    mime, _ = mimetypes.guess_type(name)
    return ResultFile(token=path.name, name=name, size=len(data), mime=mime or "application/octet-stream")


def _require_fitz():
    try:
        import fitz  # noqa: F401  (PyMuPDF)
        return True
    except ImportError:
        return False


def _require_pypdf():
    try:
        import pypdf  # noqa: F401
        return True
    except ImportError:
        return False


def _parse_ranges(spec: str, max_page: int) -> list[int]:
    """Parse '1-3,5' (1-based) into a sorted list of 0-based page indices."""
    pages: set[int] = set()
    for part in (spec or "").split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, _, b = part.partition("-")
            try:
                start, end = int(a), int(b)
            except ValueError:
                continue
            for p in range(start, end + 1):
                if 1 <= p <= max_page:
                    pages.add(p - 1)
        else:
            try:
                p = int(part)
            except ValueError:
                continue
            if 1 <= p <= max_page:
                pages.add(p - 1)
    return sorted(pages)


@register("pdf-merge")
def pdf_merge(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader, PdfWriter

    if len(files) < 2:
        return ToolResult(meta={"error": "Upload at least two PDFs to merge"})
    writer = PdfWriter()
    for src in files:
        reader = PdfReader(str(src))
        for page in reader.pages:
            writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return ToolResult(files=[_save_result(buf.getvalue(), "merged.pdf")],
                      meta={"pages": len(writer.pages)})


@register("pdf-split")
def pdf_split(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader, PdfWriter

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    reader = PdfReader(str(files[0]))
    total = len(reader.pages)
    ranges = options.get("ranges", "").strip()
    targets = _parse_ranges(ranges, total) if ranges else list(range(total))

    results: list[ResultFile] = []
    for idx in targets:
        writer = PdfWriter()
        writer.add_page(reader.pages[idx])
        buf = io.BytesIO()
        writer.write(buf)
        results.append(_save_result(buf.getvalue(), f"page_{idx + 1}.pdf"))
    return ToolResult(files=results, meta={"pages": len(results), "total": total})


@register("pdf-extract-pages")
def pdf_extract_pages(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader, PdfWriter

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    reader = PdfReader(str(files[0]))
    total = len(reader.pages)
    targets = _parse_ranges(options.get("ranges", "1"), total)
    if not targets:
        return ToolResult(meta={"error": "No valid pages selected"})
    writer = PdfWriter()
    for idx in targets:
        writer.add_page(reader.pages[idx])
    buf = io.BytesIO()
    writer.write(buf)
    return ToolResult(files=[_save_result(buf.getvalue(), "extracted.pdf")],
                      meta={"pages": len(targets)})


@register("pdf-rotate")
def pdf_rotate(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader, PdfWriter

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    angle = int(options.get("angle", 90) or 90)
    reader = PdfReader(str(files[0]))
    writer = PdfWriter()
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)
    buf = io.BytesIO()
    writer.write(buf)
    return ToolResult(files=[_save_result(buf.getvalue(), "rotated.pdf")],
                      meta={"pages": len(writer.pages), "angle": angle})


@register("pdf-metadata-viewer")
def pdf_metadata_viewer(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    try:
        reader = PdfReader(str(files[0]))
    except Exception as exc:  # noqa: BLE001 — a malformed upload is not a bug
        return ToolResult(meta={"error": f"Could not read that PDF: {exc}"})

    # pypdf returns a DocumentInformation object when the PDF carries metadata
    # and a plain {} when it does not — and {}.title raises AttributeError, so
    # a PDF with no metadata used to crash this tool outright.
    info = reader.metadata
    def field(name: str) -> str | None:
        value = getattr(info, name, None) if info is not None else None
        if value is None and isinstance(info, dict):
            value = info.get("/" + name.capitalize())
        return str(value) if value else None

    meta = {
        "pages": len(reader.pages),
        "title": field("title"),
        "author": field("author"),
        "subject": field("subject"),
        "creator": field("creator"),
        "producer": field("producer"),
        "created": field("creation_date_raw"),
        "modified": field("modification_date_raw"),
        "encrypted": reader.is_encrypted,
        "file_size_bytes": files[0].stat().st_size,
    }
    if reader.pages:
        box = reader.pages[0].mediabox
        meta["first_page_size_pt"] = f"{round(float(box.width))}x{round(float(box.height))}"
    if not any(meta[k] for k in ("title", "author", "subject", "creator", "producer")):
        meta["note"] = "This PDF carries no descriptive metadata."
    return ToolResult(meta=meta)


@register("pdf-organizer")
def pdf_organizer(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_pypdf():
        return ToolResult(meta={"error": "pypdf not installed"})
    from pypdf import PdfReader, PdfWriter

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    reader = PdfReader(str(files[0]))
    total = len(reader.pages)
    order = (options.get("order") or "").strip()
    indices = _parse_ranges(order, total) if order else list(range(total))
    if not indices:
        return ToolResult(meta={"error": "No valid pages in the requested order"})
    writer = PdfWriter()
    for idx in indices:
        writer.add_page(reader.pages[idx])
    buf = io.BytesIO()
    writer.write(buf)
    return ToolResult(files=[_save_result(buf.getvalue(), "organized.pdf")],
                      meta={"pages": len(indices)})


@register("pdf-compress")
def pdf_compress(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    original = files[0].stat().st_size
    doc = fitz.open(str(files[0]))
    out = doc.tobytes(garbage=4, deflate=True, clean=True)
    doc.close()
    rf = _save_result(out, "compressed.pdf")
    saved = max(0, original - len(out))
    return ToolResult(files=[rf], meta={
        "original_kb": round(original / 1024, 1),
        "compressed_kb": round(len(out) / 1024, 1),
        "saved_percent": round(saved / original * 100, 1) if original else 0,
    })


@register("pdf-to-jpg")
def pdf_to_jpg(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    fmt = options.get("format", "jpg")
    dpi = int(options.get("dpi", 150) or 150)
    doc = fitz.open(str(files[0]))
    results: list[ResultFile] = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap(dpi=dpi)
        data = pix.tobytes("jpg" if fmt == "jpg" else "png")
        results.append(_save_named(data, f"page_{i + 1}.{fmt}"))
    doc.close()
    return ToolResult(files=results, meta={"pages": len(results)})


@register("jpg-to-pdf")
def jpg_to_pdf(files: list[Path], text: str, options: dict) -> ToolResult:
    try:
        from PIL import Image
    except ImportError:
        return ToolResult(meta={"error": "Pillow not installed"})
    if not files:
        return ToolResult(meta={"error": "No images uploaded"})
    images = [Image.open(p).convert("RGB") for p in files]
    buf = io.BytesIO()
    images[0].save(buf, format="PDF", save_all=True, append_images=images[1:])
    return ToolResult(files=[_save_result(buf.getvalue(), "images.pdf")],
                      meta={"pages": len(images)})


@register("word-to-pdf")
def word_to_pdf(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    try:
        import docx
    except ImportError:
        return ToolResult(meta={"error": "python-docx not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    document = docx.Document(str(files[0]))
    body = "\n".join(p.text for p in document.paragraphs)
    doc = fitz.open()
    rect = fitz.Rect(56, 56, 539, 785)  # ~A4 with margins
    remaining = body
    while True:
        page = doc.new_page(width=595, height=842)
        leftover = page.insert_textbox(rect, remaining, fontsize=11, fontname="helv")
        if leftover <= 0 or not remaining.strip():
            break
        # insert_textbox returns remaining height when it fits; loop by trimming inserted text.
        # Fallback: stop to avoid infinite loop on pathological input.
        break
    out = doc.tobytes()
    doc.close()
    return ToolResult(files=[_save_result(out, "converted.pdf")],
                      meta={"note": "Text-only conversion (layout/styles not preserved)"})


@register("pdf-to-word")
def pdf_to_word(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    try:
        import docx
    except ImportError:
        return ToolResult(meta={"error": "python-docx not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    doc = fitz.open(str(files[0]))
    document = docx.Document()
    for page in doc:
        for line in page.get_text().splitlines():
            document.add_paragraph(line)
    doc.close()
    buf = io.BytesIO()
    document.save(buf)
    return ToolResult(files=[_save_named(buf.getvalue(), "converted.docx")],
                      meta={"note": "Text-only conversion (images/layout not preserved)"})


@register("pdf-unlock")
def pdf_unlock(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    doc = fitz.open(str(files[0]))
    if doc.is_encrypted:
        if not doc.authenticate(options.get("password", "") or ""):
            doc.close()
            return ToolResult(meta={"error": "Wrong password"})
    out = doc.tobytes()  # saved without encryption
    doc.close()
    return ToolResult(files=[_save_result(out, "unlocked.pdf")])


@register("pdf-protect")
def pdf_protect(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    password = options.get("password", "") or ""
    if not password:
        return ToolResult(meta={"error": "Enter a password"})
    doc = fitz.open(str(files[0]))
    perm = int(
        fitz.PDF_PERM_ACCESSIBILITY | fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY
    )
    out = doc.tobytes(
        encryption=fitz.PDF_ENCRYPT_AES_256,
        owner_pw=password, user_pw=password, permissions=perm,
    )
    doc.close()
    return ToolResult(files=[_save_result(out, "protected.pdf")])


@register("pdf-watermark")
def pdf_watermark(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    wm = options.get("text", "CONFIDENTIAL") or "CONFIDENTIAL"
    opacity = max(5, min(int(options.get("opacity", 15) or 15), 80)) / 100
    doc = fitz.open(str(files[0]))
    for page in doc:
        # Vertically centered band across the page (insert_textbox only rotates in 90° steps,
        # so we draw the watermark horizontally for reliable output).
        w, h = page.rect.width, page.rect.height
        band = fitz.Rect(0, h / 2 - 40, w, h / 2 + 40)
        page.insert_textbox(
            band, wm, fontsize=44, fontname="helv",
            color=(0.5, 0.5, 0.5), fill_opacity=opacity, align=fitz.TEXT_ALIGN_CENTER,
        )
    out = doc.tobytes()
    doc.close()
    return ToolResult(files=[_save_result(out, "watermarked.pdf")])


@register("pdf-page-numbering")
def pdf_page_numbering(files: list[Path], text: str, options: dict) -> ToolResult:
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF not installed"})
    import fitz

    if not files:
        return ToolResult(meta={"error": "No file uploaded"})
    position = options.get("position", "bottom-center")
    doc = fitz.open(str(files[0]))
    total = len(doc)
    for i, page in enumerate(doc):
        label = f"{i + 1} / {total}"
        w, h = page.rect.width, page.rect.height
        y = h - 30
        if position == "bottom-left":
            point = fitz.Point(40, y)
        elif position == "bottom-right":
            point = fitz.Point(w - 70, y)
        else:
            point = fitz.Point(w / 2 - 15, y)
        page.insert_text(point, label, fontsize=10, fontname="helv", color=(0.3, 0.3, 0.3))
    out = doc.tobytes()
    doc.close()
    return ToolResult(files=[_save_result(out, "numbered.pdf")], meta={"pages": total})


# ===========================================================================
# Shared helpers for the tools below
# ===========================================================================

def _int(options: dict, key: str, default: int) -> int:
    try:
        return int(float(str(options.get(key, default)).strip() or default))
    except (TypeError, ValueError):
        return default


def _float(options: dict, key: str, default: float) -> float:
    try:
        return float(str(options.get(key, default)).strip() or default)
    except (TypeError, ValueError):
        return default


def _flag(options: dict, key: str, default: bool = False) -> bool:
    value = options.get(key, default)
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on"}
    return bool(value)


def _stem(src: Path) -> str:
    return Path(src.name).stem.split("__", 1)[-1]


def _need_pdf(files) -> ToolResult | None:
    if not files:
        return ToolResult(meta={"error": "Upload a PDF to start."})
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF is not installed on the server."})
    return None


def _open_pdf(src: Path):
    """Open a PDF, turning a corrupt or encrypted file into a clear message."""
    import fitz

    try:
        doc = fitz.open(src)
    except Exception as exc:  # noqa: BLE001 — any malformed file lands here
        raise ValueError(f"Could not read that PDF: {exc}") from exc
    if doc.needs_pass:
        raise ValueError("That PDF is password protected. Unlock it first.")
    return doc


# ===========================================================================
# Page operations
# ===========================================================================

@register("pdf-remove-pages")
def pdf_remove_pages(files: list[Path], text: str, options: dict) -> ToolResult:
    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    spec = (text or "").strip() or str(options.get("pages", "")).strip()
    if not spec:
        return ToolResult(meta={"error": "Enter the pages to remove, e.g. 2,5-7"})
    targets = _parse_ranges(spec, doc.page_count)
    if not targets:
        return ToolResult(meta={"error": f"No valid pages in {spec!r}. This PDF has {doc.page_count}."})
    if len(targets) >= doc.page_count:
        return ToolResult(meta={"error": "That would remove every page."})
    keep = [i for i in range(doc.page_count) if i not in set(targets)]
    doc.select(keep)
    data = doc.tobytes()
    return ToolResult(files=[_save_result(data, f"trimmed_{_stem(files[0])}.pdf")],
                      meta={"pages_removed": len(targets), "pages_left": len(keep)})


@register("pdf-insert-pages")
def pdf_insert_pages(files: list[Path], text: str, options: dict) -> ToolResult:
    """Insert one PDF into another at a chosen page."""
    guard = _need_pdf(files)
    if guard:
        return guard
    if len(files) < 2:
        return ToolResult(meta={"error": "Upload two PDFs: the original first, then the one to insert."})
    try:
        base, extra = _open_pdf(files[0]), _open_pdf(files[1])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    position = _int(options, "after_page", 0)
    # 0 means "before page 1"; anything else lands after the numbered page.
    at = max(0, min(position, base.page_count))
    base.insert_pdf(extra, start_at=at)
    return ToolResult(files=[_save_result(base.tobytes(), f"merged_{_stem(files[0])}.pdf")],
                      meta={"inserted_pages": extra.page_count,
                            "inserted_after_page": at,
                            "total_pages": base.page_count})


@register("pdf-alternate-pages")
def pdf_alternate_pages(files: list[Path], text: str, options: dict) -> ToolResult:
    """Interleave two PDFs — for a scanner that does odd and even sides separately."""
    guard = _need_pdf(files)
    if guard:
        return guard
    if len(files) < 2:
        return ToolResult(meta={"error": "Upload both halves: odd pages first, then even."})
    import fitz

    try:
        first, second = _open_pdf(files[0]), _open_pdf(files[1])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    if _flag(options, "reverse_second"):
        # A duplex scanner that feeds the second pass backwards produces the
        # even pages in reverse; this is the fix for that.
        second.select(list(range(second.page_count - 1, -1, -1)))
    out = fitz.open()
    for index in range(max(first.page_count, second.page_count)):
        if index < first.page_count:
            out.insert_pdf(first, from_page=index, to_page=index)
        if index < second.page_count:
            out.insert_pdf(second, from_page=index, to_page=index)
    return ToolResult(files=[_save_result(out.tobytes(), "interleaved.pdf")],
                      meta={"pages": out.page_count,
                            "from_first": first.page_count, "from_second": second.page_count})


@register("pdf-resize-pages")
def pdf_resize_pages(files: list[Path], text: str, options: dict) -> ToolResult:
    """Scale every page onto a standard paper size."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    paper = str(options.get("size", "a4"))
    try:
        target = fitz.paper_rect(paper if not _flag(options, "landscape") else f"{paper}-l")
    except Exception:  # noqa: BLE001 — unknown paper name
        return ToolResult(meta={"error": f"Unknown paper size {paper!r}."})
    margin = max(0, min(_int(options, "margin", 0), 100))
    out = fitz.open()
    for page in doc:
        new_page = out.new_page(width=target.width, height=target.height)
        box = target + (margin, margin, -margin, -margin)
        # show_pdf_page scales the source page into the box, keeping its ratio,
        # so nothing is stretched or cropped.
        new_page.show_pdf_page(box, doc, page.number, keep_proportion=True)
    return ToolResult(files=[_save_result(out.tobytes(), f"{paper}_{_stem(files[0])}.pdf")],
                      meta={"pages": out.page_count, "size": paper,
                            "dimensions_pt": f"{round(target.width)}x{round(target.height)}"})


@register("pdf-crop")
def pdf_crop(files: list[Path], text: str, options: dict) -> ToolResult:
    """Trim the margins off every page, by percentage or automatically."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    auto = _flag(options, "auto", False)
    percents = {side: max(0.0, min(_float(options, side, 0), 45))
                for side in ("top", "bottom", "left", "right")}
    cropped = 0
    for page in doc:
        rect = page.rect
        if auto:
            # The bounding box of everything drawn on the page, plus a small
            # margin — this is what "crop to content" means.
            content = page.get_bboxlog()
            boxes = [fitz.Rect(b[1]) for b in content if b[1]]
            if not boxes:
                continue
            box = boxes[0]
            for b in boxes[1:]:
                box |= b
            pad = _float(options, "auto_margin", 6)
            box = fitz.Rect(max(rect.x0, box.x0 - pad), max(rect.y0, box.y0 - pad),
                            min(rect.x1, box.x1 + pad), min(rect.y1, box.y1 + pad))
        else:
            box = fitz.Rect(rect.x0 + rect.width * percents["left"] / 100,
                            rect.y0 + rect.height * percents["top"] / 100,
                            rect.x1 - rect.width * percents["right"] / 100,
                            rect.y1 - rect.height * percents["bottom"] / 100)
        if box.width > 10 and box.height > 10:
            page.set_cropbox(box)
            cropped += 1
    return ToolResult(files=[_save_result(doc.tobytes(), f"cropped_{_stem(files[0])}.pdf")],
                      meta={"pages_cropped": cropped, "mode": "auto" if auto else "percentage"})


@register("pdf-grayscale")
def pdf_grayscale(files: list[Path], text: str, options: dict) -> ToolResult:
    """Convert to greyscale — cheaper to print, and often a much smaller file."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    dpi = max(72, min(_int(options, "dpi", 150), 400))
    out = fitz.open()
    for page in doc:
        # Rendering to a greyscale bitmap is the only way to guarantee no colour
        # survives; keeping the vector content would leave coloured text behind.
        pix = page.get_pixmap(dpi=dpi, colorspace=fitz.csGRAY)
        new_page = out.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(page.rect, pixmap=pix)
    data = out.tobytes(deflate=True, garbage=4)
    return ToolResult(files=[_save_result(data, f"grayscale_{_stem(files[0])}.pdf")],
                      meta={"pages": out.page_count, "dpi": dpi,
                            "before_bytes": files[0].stat().st_size, "after_bytes": len(data),
                            "note": "Pages become images, so the text is no longer selectable."})


@register("pdf-flatten")
def pdf_flatten(files: list[Path], text: str, options: dict) -> ToolResult:
    """Bake annotations and form fields into the page so nothing stays editable."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    annotations = sum(len(list(page.annots() or [])) for page in doc)
    widgets = sum(len(list(page.widgets() or [])) for page in doc)
    if _flag(options, "rasterize"):
        dpi = max(72, min(_int(options, "dpi", 150), 400))
        out = fitz.open()
        for page in doc:
            pix = page.get_pixmap(dpi=dpi)
            new_page = out.new_page(width=page.rect.width, height=page.rect.height)
            new_page.insert_image(page.rect, pixmap=pix)
        data = out.tobytes(deflate=True, garbage=4)
        note = "Rasterised — text is no longer selectable, but nothing can be edited back out."
    else:
        # bake() draws the appearance streams into the page content and drops
        # the interactive objects, keeping the text real.
        try:
            doc.bake()
        except Exception:  # noqa: BLE001 — older PyMuPDF without bake()
            for page in doc:
                for annot in list(page.annots() or []):
                    page.delete_annot(annot)
        data = doc.tobytes(deflate=True, garbage=4)
        note = "Annotations and fields baked in; text stays selectable."
    return ToolResult(files=[_save_result(data, f"flat_{_stem(files[0])}.pdf")],
                      meta={"annotations_flattened": annotations, "form_fields_flattened": widgets,
                            "note": note})


@register("pdf-repair")
def pdf_repair(files: list[Path], text: str, options: dict) -> ToolResult:
    """Rebuild a damaged PDF's structure, keeping whatever content survives."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    src = files[0]
    try:
        doc = fitz.open(src)
    except Exception as exc:  # noqa: BLE001
        return ToolResult(meta={"error": f"That file is too damaged to open at all: {exc}"})
    if doc.needs_pass:
        return ToolResult(meta={"error": "That PDF is password protected. Unlock it first."})
    readable, broken = 0, []
    for page in doc:
        try:
            page.get_text()
            readable += 1
        except Exception:  # noqa: BLE001 — a single unreadable page is the point here
            broken.append(page.number + 1)
    # garbage=4 rebuilds the cross-reference table and drops orphaned objects,
    # which is what fixes most "cannot open" files.
    data = doc.tobytes(garbage=4, deflate=True, clean=True)
    return ToolResult(files=[_save_result(data, f"repaired_{_stem(src)}.pdf")],
                      meta={"pages": doc.page_count, "readable_pages": readable,
                            "problem_pages": broken or "none",
                            "before_bytes": src.stat().st_size, "after_bytes": len(data)})


@register("pdf-optimize-web")
def pdf_optimize_web(files: list[Path], text: str, options: dict) -> ToolResult:
    """Linearise for fast web viewing, and optionally downsample the images.

    Linearising rewrites the file so a browser can show page one before the rest
    has downloaded — which is the difference between a PDF that opens instantly
    and one that shows a spinner.
    """
    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    before = files[0].stat().st_size
    if _flag(options, "downsample_images", True):
        limit = max(72, min(_int(options, "image_dpi", 150), 400))
        try:
            doc.rewrite_images(dpi_target=limit, quality=_int(options, "quality", 80))
        except Exception:  # noqa: BLE001 — older PyMuPDF without rewrite_images
            pass
    data = doc.tobytes(linear=True, garbage=4, deflate=True, clean=True)
    return ToolResult(files=[_save_result(data, f"web_{_stem(files[0])}.pdf")],
                      meta={"pages": doc.page_count,
                            "before_bytes": before, "after_bytes": len(data),
                            "saved_percent": round(100 - len(data) / max(1, before) * 100, 1),
                            "linearized": True})


@register("pdf-enhance-scan")
def pdf_enhance_scan(files: list[Path], text: str, options: dict) -> ToolResult:
    """Whiten and sharpen the pages of a scanned document."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import io as _io

    import fitz
    import numpy as np
    from PIL import Image, ImageFilter, ImageOps

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    dpi = max(100, min(_int(options, "dpi", 200), 400))
    out = fitz.open()
    for page in doc:
        pix = page.get_pixmap(dpi=dpi)
        img = Image.open(_io.BytesIO(pix.tobytes("png"))).convert("L")
        if _flag(options, "whiten", True):
            # Dividing by a blurred copy removes uneven lighting across the page.
            background = img.filter(ImageFilter.GaussianBlur(radius=max(5, img.width // 40)))
            a = np.asarray(img, dtype=float)
            b = np.asarray(background, dtype=float)
            img = Image.fromarray(np.clip(a / np.maximum(b, 1) * 255, 0, 255).astype("uint8"))
            img = ImageOps.autocontrast(img, cutoff=(1, 20))
        if _flag(options, "sharpen", True):
            img = img.filter(ImageFilter.UnsharpMask(radius=1.5, percent=130, threshold=3))
        buf = _io.BytesIO()
        img.save(buf, "JPEG", quality=_int(options, "quality", 80))
        new_page = out.new_page(width=page.rect.width, height=page.rect.height)
        new_page.insert_image(page.rect, stream=buf.getvalue())
    data = out.tobytes(deflate=True, garbage=4)
    return ToolResult(files=[_save_result(data, f"enhanced_{_stem(files[0])}.pdf")],
                      meta={"pages": out.page_count, "dpi": dpi,
                            "before_bytes": files[0].stat().st_size, "after_bytes": len(data)})


# ===========================================================================
# Content: add, redact, sign
# ===========================================================================

@register("pdf-add-text")
def pdf_add_text(files: list[Path], text: str, options: dict) -> ToolResult:
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    message = (text or "").strip() or str(options.get("message", "")).strip()
    if not message:
        return ToolResult(meta={"error": "Type the text to add."})
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    spec = str(options.get("pages", "")).strip()
    targets = _parse_ranges(spec, doc.page_count) if spec else list(range(doc.page_count))
    if not targets:
        return ToolResult(meta={"error": f"No valid pages in {spec!r}."})
    size = max(6, min(_int(options, "font_size", 14), 96))
    colour = str(options.get("color", "#000000")).lstrip("#")
    try:
        rgb = tuple(int(colour[i:i + 2], 16) / 255 for i in (0, 2, 4))
    except (ValueError, IndexError):
        rgb = (0, 0, 0)
    position = str(options.get("position", "top-left"))
    for index in targets:
        page = doc[index]
        rect = page.rect
        margin = 40
        x = {"top-left": margin, "bottom-left": margin}.get(position)
        if x is None:
            width = fitz.get_text_length(message, fontname="helv", fontsize=size)
            x = (rect.width - width) / 2 if "center" in position else rect.width - width - margin
        y = margin + size if "top" in position else rect.height - margin
        page.insert_text((x + _int(options, "offset_x", 0), y + _int(options, "offset_y", 0)),
                         message, fontsize=size, color=rgb, fontname="helv")
    return ToolResult(files=[_save_result(doc.tobytes(), f"text_{_stem(files[0])}.pdf")],
                      meta={"pages_updated": len(targets)})


@register("pdf-add-image")
def pdf_add_image(files: list[Path], text: str, options: dict) -> ToolResult:
    """Stamp a logo or picture onto chosen pages."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    if len(files) < 2:
        return ToolResult(meta={"error": "Upload the PDF first, then the image to place on it."})
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    image_bytes = files[1].read_bytes()
    try:
        picture = fitz.Pixmap(image_bytes)
    except Exception:  # noqa: BLE001
        return ToolResult(meta={"error": "The second file is not a readable image."})
    spec = str(options.get("pages", "")).strip()
    targets = _parse_ranges(spec, doc.page_count) if spec else list(range(doc.page_count))
    width_percent = max(1, min(_float(options, "width_percent", 25), 100)) / 100
    margin = _int(options, "margin", 24)
    position = str(options.get("position", "top-right"))
    for index in targets:
        page = doc[index]
        rect = page.rect
        w = rect.width * width_percent
        h = w * picture.height / max(1, picture.width)
        x0 = margin if "left" in position else (rect.width - w) / 2 if "center" in position else rect.width - w - margin
        y0 = margin if "top" in position else (rect.height - h) / 2 if "middle" in position else rect.height - h - margin
        page.insert_image(fitz.Rect(x0, y0, x0 + w, y0 + h), stream=image_bytes,
                          overlay=not _flag(options, "behind_text"))
    return ToolResult(files=[_save_result(doc.tobytes(), f"stamped_{_stem(files[0])}.pdf")],
                      meta={"pages_updated": len(targets)})


@register("pdf-header-footer")
def pdf_header_footer(files: list[Path], text: str, options: dict) -> ToolResult:
    """Running header and footer, with placeholders for the page numbers."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    header = str(options.get("header", "")).strip()
    footer = str(options.get("footer", "")).strip() or (text or "").strip()
    if not header and not footer:
        return ToolResult(meta={"error": "Enter a header, a footer, or both."})
    size = max(6, min(_int(options, "font_size", 9), 24))
    margin = _int(options, "margin", 28)
    skip_first = _flag(options, "skip_first_page")
    for page in doc:
        if skip_first and page.number == 0:
            continue
        for content, at_top in ((header, True), (footer, False)):
            if not content:
                continue
            # {page} and {total} are what people expect to be able to write.
            rendered = (content.replace("{page}", str(page.number + 1))
                               .replace("{total}", str(doc.page_count)))
            width = fitz.get_text_length(rendered, fontname="helv", fontsize=size)
            align = str(options.get("align", "center"))
            x = (margin if align == "left" else
                 page.rect.width - width - margin if align == "right"
                 else (page.rect.width - width) / 2)
            y = margin if at_top else page.rect.height - margin / 2
            page.insert_text((x, y), rendered, fontsize=size, fontname="helv",
                             color=(0.35, 0.35, 0.35))
    return ToolResult(files=[_save_result(doc.tobytes(), f"headed_{_stem(files[0])}.pdf")],
                      meta={"pages": doc.page_count,
                            "note": "Use {page} and {total} for the numbers."})


@register("pdf-redact")
def pdf_redact(files: list[Path], text: str, options: dict) -> ToolResult:
    """True redaction: the text is deleted, not covered over.

    A black rectangle drawn on top still has the words underneath, and anyone
    can select and copy them. This removes the content itself, which is the
    only version that survives someone opening the file in an editor.
    """
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    terms = [t.strip() for t in (text or "").split("\n") if t.strip()]
    if not terms:
        return ToolResult(meta={"error": "Enter the words or phrases to redact, one per line."})
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    ignore_case = _flag(options, "ignore_case", True)
    found = 0
    for page in doc:
        for term in terms:
            flags = fitz.TEXT_DEHYPHENATE
            hits = page.search_for(term, flags=flags)
            if ignore_case and not hits:
                hits = page.search_for(term.lower(), flags=flags) or page.search_for(term.upper(), flags=flags)
            for rect in hits:
                page.add_redact_annot(rect, fill=(0, 0, 0))
                found += 1
        page.apply_redactions()
    if not found:
        return ToolResult(meta={"error": "None of those terms were found in this PDF."})
    data = doc.tobytes(garbage=4, deflate=True)
    return ToolResult(files=[_save_result(data, f"redacted_{_stem(files[0])}.pdf")],
                      meta={"redactions": found, "terms": terms,
                            "note": "The text is removed from the file, not just covered."})


@register("pdf-sign")
def pdf_sign(files: list[Path], text: str, options: dict) -> ToolResult:
    """Place a signature image on the page.

    A visible signature, not a cryptographic one — it proves nothing about who
    signed or whether the document changed afterwards. For that you need a
    certificate-based signature from a provider.
    """
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    if len(files) < 2:
        return ToolResult(meta={"error": "Upload the PDF first, then your signature image."})
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    signature = files[1].read_bytes()
    try:
        pix = fitz.Pixmap(signature)
    except Exception:  # noqa: BLE001
        return ToolResult(meta={"error": "The second file is not a readable image."})
    page_number = _int(options, "page", 0)
    index = doc.page_count - 1 if page_number <= 0 else min(page_number, doc.page_count) - 1
    page = doc[index]
    width = max(40, min(_int(options, "width", 160), int(page.rect.width)))
    height = width * pix.height / max(1, pix.width)
    x = _int(options, "x", 0) or int(page.rect.width - width - 60)
    y = _int(options, "y", 0) or int(page.rect.height - height - 90)
    page.insert_image(fitz.Rect(x, y, x + width, y + height), stream=signature, overlay=True)
    if _flag(options, "add_date", True):
        from datetime import date
        page.insert_text((x, y + height + 12),
                         f"Signed {date.today().isoformat()}", fontsize=8,
                         fontname="helv", color=(0.35, 0.35, 0.35))
    if _flag(options, "lock", True):
        # Flattening afterwards stops the image being dragged off in an editor.
        try:
            doc.bake()
        except Exception:  # noqa: BLE001
            pass
    return ToolResult(files=[_save_result(doc.tobytes(), f"signed_{_stem(files[0])}.pdf")],
                      meta={"signed_page": index + 1,
                            "warning": "A visible signature only — it is not a legally "
                                       "verifiable digital signature."})


@register("pdf-create-blank")
def pdf_create_blank(files: list[Path], text: str, options: dict) -> ToolResult:
    """A blank PDF — lined, gridded, dotted or plain."""
    import fitz

    pages = max(1, min(_int(options, "pages", 1), 200))
    paper = str(options.get("size", "a4"))
    try:
        rect = fitz.paper_rect(paper if not _flag(options, "landscape") else f"{paper}-l")
    except Exception:  # noqa: BLE001
        return ToolResult(meta={"error": f"Unknown paper size {paper!r}."})
    style = str(options.get("style", "blank"))
    spacing = max(6, min(_int(options, "spacing", 24), 100))
    doc = fitz.open()
    grey = (0.75, 0.75, 0.8)
    for _ in range(pages):
        page = doc.new_page(width=rect.width, height=rect.height)
        if style == "lined":
            y = spacing * 2
            while y < rect.height - spacing:
                page.draw_line((40, y), (rect.width - 40, y), color=grey, width=0.5)
                y += spacing
        elif style == "grid":
            for x in range(40, int(rect.width) - 39, spacing):
                page.draw_line((x, 40), (x, rect.height - 40), color=grey, width=0.4)
            for y in range(40, int(rect.height) - 39, spacing):
                page.draw_line((40, y), (rect.width - 40, y), color=grey, width=0.4)
        elif style == "dotted":
            for x in range(40, int(rect.width) - 39, spacing):
                for y in range(40, int(rect.height) - 39, spacing):
                    page.draw_circle((x, y), 0.7, color=grey, fill=grey)
    return ToolResult(files=[_save_result(doc.tobytes(), f"blank-{style}-{paper}.pdf")],
                      meta={"pages": pages, "size": paper, "style": style})


@register("pdf-bates-numbering")
def pdf_bates_numbering(files: list[Path], text: str, options: dict) -> ToolResult:
    """Sequential legal numbering across every uploaded file.

    Bates numbers run continuously through a whole production, not per file, so
    uploading several PDFs numbers them as one set in the order given.
    """
    guard = _need_pdf(files)
    if guard:
        return guard
    prefix = str(options.get("prefix", "")).strip()
    suffix = str(options.get("suffix", "")).strip()
    digits = max(1, min(_int(options, "digits", 6), 12))
    counter = max(0, _int(options, "start", 1))
    size = max(6, min(_int(options, "font_size", 10), 24))
    position = str(options.get("position", "bottom-right"))
    results, first, last = [], None, None
    for src in files:
        try:
            doc = _open_pdf(src)
        except ValueError as e:
            return ToolResult(meta={"error": f"{_stem(src)}: {e}"})
        import fitz

        for page in doc:
            label = f"{prefix}{str(counter).zfill(digits)}{suffix}"
            first = first or label
            last = label
            width = fitz.get_text_length(label, fontname="helv", fontsize=size)
            x = 36 if "left" in position else page.rect.width - width - 36
            y = 30 if "top" in position else page.rect.height - 24
            page.insert_text((x, y), label, fontsize=size, fontname="helv", color=(0, 0, 0))
            counter += 1
        results.append(_save_result(doc.tobytes(), f"bates_{_stem(src)}.pdf"))
    return ToolResult(files=results, meta={
        "files": len(results), "pages_numbered": counter - _int(options, "start", 1),
        "first_number": first, "last_number": last,
    })


@register("pdf-remove-metadata")
def pdf_remove_metadata(files: list[Path], text: str, options: dict) -> ToolResult:
    """Strip author, software and timestamps from a PDF."""
    guard = _need_pdf(files)
    if guard:
        return guard
    results, report = [], []
    for src in files:
        try:
            doc = _open_pdf(src)
        except ValueError as e:
            return ToolResult(meta={"error": f"{_stem(src)}: {e}"})
        had = {k: v for k, v in (doc.metadata or {}).items() if v}
        doc.set_metadata({})
        try:
            # The XMP packet is a second, separate copy of the metadata — clearing
            # only the document info dictionary leaves the author name in there.
            doc.del_xml_metadata()
        except Exception:  # noqa: BLE001
            pass
        data = doc.tobytes(garbage=4, deflate=True, clean=True)
        results.append(_save_result(data, f"clean_{_stem(src)}.pdf"))
        report.append({"file": _stem(src), "removed": sorted(had),
                       "author_was": had.get("author") or None,
                       "producer_was": had.get("producer") or None})
    return ToolResult(files=results, meta={"count": len(results), "files": report})


@register("pdf-compare")
def pdf_compare(files: list[Path], text: str, options: dict) -> ToolResult:
    """Compare two PDFs page by page, on their text."""
    import difflib

    guard = _need_pdf(files)
    if guard:
        return guard
    if len(files) < 2:
        return ToolResult(meta={"error": "Upload both PDFs to compare."})
    try:
        left, right = _open_pdf(files[0]), _open_pdf(files[1])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    ignore_space = _flag(options, "ignore_whitespace", True)

    def words(doc, index):
        raw = doc[index].get_text() if index < doc.page_count else ""
        return raw.split() if ignore_space else raw.splitlines()

    pages, changed = [], 0
    diff_lines = []
    for index in range(max(left.page_count, right.page_count)):
        a, b = words(left, index), words(right, index)
        ratio = difflib.SequenceMatcher(None, a, b).ratio() if (a or b) else 1.0
        if ratio < 0.999:
            changed += 1
            diff_lines.append(f"--- page {index + 1} ---")
            diff_lines += [ln for ln in difflib.unified_diff(a, b, lineterm="", n=1)
                           if ln.startswith(("+", "-")) and not ln.startswith(("+++", "---"))][:40]
        pages.append({"page": index + 1, "similarity_percent": round(ratio * 100, 1),
                      "identical": ratio >= 0.999})
    return ToolResult(text="\n".join(diff_lines) or "The two documents match.", meta={
        "identical": changed == 0,
        "pages_compared": len(pages), "pages_changed": changed,
        "first_pages": left.page_count, "second_pages": right.page_count,
        "pages": pages[:60],
        "note": "Text comparison — a change to an image or to layout alone is not detected.",
    })


@register("pdf-to-base64")
def pdf_to_base64(files: list[Path], text: str, options: dict) -> ToolResult:
    import base64

    if not files:
        return ToolResult(meta={"error": "Upload a PDF."})
    data = files[0].read_bytes()
    # 8 MB of PDF becomes ~11 MB of Base64 in the response body; past that the
    # browser struggles to hold it in a textarea.
    if len(data) > 8 * 1024 * 1024:
        return ToolResult(meta={"error": "That PDF is over 8 MB — too large to encode inline."})
    encoded = base64.b64encode(data).decode()
    if _flag(options, "data_uri", True):
        encoded = "data:application/pdf;base64," + encoded
    return ToolResult(text=encoded,
                      files=[_save_result(encoded.encode(), f"{_stem(files[0])}.b64.txt")],
                      meta={"file_bytes": len(data), "encoded_bytes": len(encoded),
                            "growth_percent": round(len(encoded) / max(1, len(data)) * 100 - 100, 1)})


# ===========================================================================
# Extraction
# ===========================================================================

@register("pdf-extract-text")
def pdf_extract_text(files: list[Path], text: str, options: dict) -> ToolResult:
    """Pull the text out, as plain text, Markdown or JSON per page."""
    import json as _json

    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    layout = _flag(options, "keep_layout")
    pages = []
    for page in doc:
        # "blocks" keeps the reading order of columns; plain text runs them
        # together, which scrambles a two-column page.
        content = page.get_text("text" if not layout else "blocks")
        if layout:
            content = "\n".join(b[4] for b in sorted(content, key=lambda b: (round(b[1]), b[0]))
                                if len(b) > 4 and isinstance(b[4], str))
        pages.append(content.strip())
    fmt = str(options.get("format", "txt"))
    if fmt == "json":
        body = _json.dumps([{"page": i + 1, "text": t} for i, t in enumerate(pages)],
                           indent=2, ensure_ascii=False)
        ext = "json"
    elif fmt == "markdown":
        body = "\n\n".join(f"## Page {i + 1}\n\n{t}" for i, t in enumerate(pages) if t)
        ext = "md"
    else:
        separator = "\n\n" + ("-" * 40) + "\n\n" if _flag(options, "page_breaks", True) else "\n\n"
        body = separator.join(t for t in pages if t)
        ext = "txt"
    words = sum(len(t.split()) for t in pages)
    empty = [i + 1 for i, t in enumerate(pages) if not t]
    return ToolResult(
        text=body[:20000] + ("\n\n…truncated in the preview; the download has everything."
                             if len(body) > 20000 else ""),
        files=[_save_named(body.encode("utf-8"), f"{_stem(files[0])}.{ext}")],
        meta={"pages": doc.page_count, "words": words, "characters": len(body),
              "pages_with_no_text": empty or "none",
              "note": ("Some pages hold no text — they are probably scans, which need OCR."
                       if empty else "Text extracted from every page.")})


@register("pdf-to-html")
def pdf_to_html(files: list[Path], text: str, options: dict) -> ToolResult:
    guard = _need_pdf(files)
    if guard:
        return guard
    from html import escape

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    if _flag(options, "preserve_layout"):
        # PyMuPDF's own HTML keeps absolute positions — it looks like the PDF,
        # but the markup is not something you would hand-edit afterwards.
        body = "".join(page.get_text("html") for page in doc)
    else:
        parts = []
        for page in doc:
            parts.append(f'<section class="page" id="page-{page.number + 1}">')
            for block in page.get_text("blocks"):
                content = block[4] if len(block) > 4 and isinstance(block[4], str) else ""
                content = content.strip()
                if not content:
                    continue
                # A short block in a larger font is almost always a heading.
                lines = content.split("\n")
                if len(lines) == 1 and len(content) < 80:
                    parts.append(f"<h2>{escape(content)}</h2>")
                else:
                    parts.append("<p>" + "<br>".join(escape(ln) for ln in lines) + "</p>")
            parts.append("</section>")
        body = "\n".join(parts)
    title = escape(_stem(files[0]))
    html = (f'<!doctype html>\n<html lang="en">\n<head>\n<meta charset="utf-8">\n'
            f'<meta name="viewport" content="width=device-width, initial-scale=1">\n'
            f"<title>{title}</title>\n<style>\n"
            "body{font:16px/1.6 system-ui,sans-serif;max-width:820px;margin:2rem auto;padding:0 1rem}\n"
            ".page{margin-bottom:3rem}h2{margin:1.5rem 0 .5rem;font-size:1.25rem}\n"
            "</style>\n</head>\n<body>\n" + body + "\n</body>\n</html>")
    return ToolResult(text=html[:20000],
                      files=[_save_named(html.encode("utf-8"), f"{_stem(files[0])}.html")],
                      meta={"pages": doc.page_count, "bytes": len(html)})


@register("pdf-extract-images")
def pdf_extract_images(files: list[Path], text: str, options: dict) -> ToolResult:
    """Pull out the embedded images at their original resolution."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    minimum = max(0, _int(options, "min_size", 100))
    limit = max(1, min(_int(options, "max_images", 100), 500))
    results, seen, skipped = [], set(), 0
    for page in doc:
        for info in page.get_images(full=True):
            xref = info[0]
            if xref in seen:
                continue  # a logo repeated on every page is one image, not thirty
            seen.add(xref)
            try:
                extracted = doc.extract_image(xref)
            except Exception:  # noqa: BLE001 — an unsupported codec is not a failure
                continue
            width, height = extracted.get("width", 0), extracted.get("height", 0)
            if width < minimum or height < minimum:
                skipped += 1
                continue
            ext = extracted.get("ext", "png")
            results.append(_save_named(extracted["image"],
                                       f"p{page.number + 1}_img{len(results) + 1}.{ext}"))
            if len(results) >= limit:
                break
        if len(results) >= limit:
            break
    if not results:
        return ToolResult(meta={
            "error": "No embedded images found. A scanned page is one big image — "
                     "use PDF to JPG for that instead."
        })
    return ToolResult(files=results, meta={
        "images": len(results), "skipped_below_minimum": skipped,
        "note": "Duplicates across pages are returned once.",
    })


@register("pdf-extract-tables")
def pdf_extract_tables(files: list[Path], text: str, options: dict) -> ToolResult:
    """Find tables and export them as CSV."""
    import csv as _csv
    import io as _io

    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    results, summary = [], []
    combined = []
    for page in doc:
        try:
            found = page.find_tables()
        except Exception:  # noqa: BLE001 — older PyMuPDF without find_tables
            return ToolResult(meta={"error": "This server's PyMuPDF is too old to find tables."})
        for index, table in enumerate(found.tables, start=1):
            rows = table.extract()
            rows = [[("" if cell is None else str(cell).strip()) for cell in row] for row in rows]
            rows = [r for r in rows if any(r)]
            if len(rows) < 2:
                continue  # a single row is a heading, not a table
            buf = _io.StringIO()
            _csv.writer(buf, lineterminator="\n").writerows(rows)
            name = f"page{page.number + 1}_table{index}.csv"
            results.append(_save_named(buf.getvalue().encode("utf-8"), name))
            combined.append(f"--- page {page.number + 1}, table {index} ---\n" + buf.getvalue())
            summary.append({"page": page.number + 1, "table": index,
                            "rows": len(rows), "columns": max(len(r) for r in rows)})
    if not results:
        return ToolResult(meta={
            "error": "No tables detected. Tables drawn without ruling lines, and tables "
                     "inside a scanned image, cannot be found this way."
        })
    return ToolResult(files=results, text="\n\n".join(combined)[:20000],
                      meta={"tables": len(results), "detail": summary})


@register("pdf-extract-attachments")
def pdf_extract_attachments(files: list[Path], text: str, options: dict) -> ToolResult:
    """Files embedded inside the PDF — invoices often carry an XML copy."""
    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    results, rows = [], []
    for index in range(doc.embfile_count()):
        info = doc.embfile_info(index)
        data = doc.embfile_get(index)
        name = info.get("filename") or info.get("name") or f"attachment_{index + 1}"
        safe = "".join(c for c in name if c.isalnum() or c in "._- ") or f"attachment_{index + 1}"
        results.append(_save_named(data, safe))
        rows.append({"name": name, "bytes": len(data), "description": info.get("desc") or None})
    if not results:
        return ToolResult(meta={"error": "This PDF has no embedded attachments."})
    return ToolResult(files=results, meta={"attachments": len(results), "files": rows})


@register("pdf-summarize")
def pdf_summarize(files: list[Path], text: str, options: dict) -> ToolResult:
    """Extractive summary — keeps the document's own highest-scoring sentences.

    Nothing is generated, so the summary cannot state something the PDF does
    not. Every sentence below appears verbatim in the source.
    """
    import re as _re
    from collections import Counter

    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    body = " ".join(page.get_text() for page in doc)
    body = _re.sub(r"\s+", " ", body).strip()
    if len(body.split()) < 60:
        return ToolResult(meta={"error": "Not enough text to summarise. Scanned PDFs need OCR first."})
    sentences = [s.strip() for s in _re.split(r"(?<=[.!?])\s+", body) if len(s.split()) > 4]
    if len(sentences) < 4:
        return ToolResult(meta={"error": "Not enough complete sentences to summarise."})
    want = max(1, min(_int(options, "sentences", 5), len(sentences)))
    stop = {"the","a","an","and","or","of","to","in","on","for","with","is","are","was","were",
            "be","it","this","that","as","at","by","from","but","not","have","has","which"}
    freq = Counter(w.lower() for w in _re.findall(r"[A-Za-z']+", body)
                   if w.lower() not in stop and len(w) > 2)
    peak = max(freq.values()) if freq else 1
    scored = []
    for index, sentence in enumerate(sentences):
        words = [w.lower() for w in _re.findall(r"[A-Za-z']+", sentence)]
        if not words:
            continue
        # Divided by length, so a long sentence does not win on word count alone.
        scored.append((sum(freq.get(w, 0) / peak for w in words) / len(words), index, sentence))
    picked = sorted(sorted(scored, reverse=True)[:want], key=lambda t: t[1])
    summary = " ".join(s for _, _, s in picked)
    return ToolResult(text=summary,
                      files=[_save_named(summary.encode("utf-8"), f"summary_{_stem(files[0])}.txt")],
                      meta={"pages": doc.page_count,
                            "source_words": len(body.split()),
                            "summary_words": len(summary.split()),
                            "reduction_percent": round(100 - len(summary.split()) /
                                                       max(1, len(body.split())) * 100, 1),
                            "keywords": [w for w, _ in freq.most_common(10)],
                            "note": "Extractive — every sentence is taken verbatim from the PDF."})


# ===========================================================================
# Conversions into PDF
# ===========================================================================

def _story_to_pdf(html: str, css: str = "", paper: str = "a4", margin: int = 50) -> bytes:
    """Render HTML into a paginated PDF with PyMuPDF's Story engine."""
    import io as _io

    import fitz

    buf = _io.BytesIO()
    story = fitz.Story(html=html, user_css=css or None)
    writer = fitz.DocumentWriter(buf)
    mediabox = fitz.paper_rect(paper)
    where = mediabox + (margin, margin, -margin, -margin)
    more, pages = 1, 0
    while more:
        device = writer.begin_page(mediabox)
        more, _ = story.place(where)
        story.draw(device)
        writer.end_page()
        pages += 1
        # A runaway layout would otherwise loop until the box runs out of memory.
        if pages > 2000:
            break
    writer.close()
    return buf.getvalue()


@register("html-to-pdf")
def html_to_pdf(files: list[Path], text: str, options: dict) -> ToolResult:
    """Render pasted HTML into a PDF.

    Pasted markup only — this server never fetches a URL on your behalf, which
    is what stops it being used to reach machines behind a firewall.
    """
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF is not installed on the server."})
    html = (text or "").strip()
    if not html and files:
        html = files[0].read_text(errors="replace")
    if not html:
        return ToolResult(meta={"error": "Paste your HTML, or upload an .html file."})
    if html.lower().startswith(("http://", "https://")):
        return ToolResult(meta={
            "error": "Paste the page's HTML, not its URL. Open the page, view source, "
                     "and copy it in — this server does not fetch external pages."
        })
    paper = str(options.get("size", "a4"))
    try:
        data = _story_to_pdf(html, str(options.get("css", "")), paper,
                             max(0, min(_int(options, "margin", 50), 200)))
    except Exception as exc:  # noqa: BLE001 — malformed HTML, not a bug
        return ToolResult(meta={"error": f"Could not lay that out: {exc}"})
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    return ToolResult(files=[_save_result(data, "converted.pdf")],
                      meta={"pages": doc.page_count, "bytes": len(data), "size": paper,
                            "note": "Scripts and external stylesheets are ignored — "
                                    "inline your CSS in the box below."})


@register("text-to-pdf")
def text_to_pdf(files: list[Path], text: str, options: dict) -> ToolResult:
    """Turn plain text or Markdown into a formatted PDF."""
    from html import escape

    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF is not installed on the server."})
    body = (text or "").strip()
    if not body and files:
        body = files[0].read_text(errors="replace")
    if not body:
        return ToolResult(meta={"error": "Paste your text, or upload a .txt or .md file."})

    if _flag(options, "markdown", True):
        import re as _re
        lines = []
        for raw in body.split("\n"):
            line = raw.rstrip()
            heading = _re.match(r"^(#{1,6})\s+(.*)$", line)
            if heading:
                level = len(heading.group(1))
                lines.append(f"<h{level}>{escape(heading.group(2))}</h{level}>")
                continue
            bullet = _re.match(r"^\s*[-*+]\s+(.*)$", line)
            if bullet:
                lines.append(f"<li>{escape(bullet.group(1))}</li>")
                continue
            if not line.strip():
                lines.append("")
                continue
            escaped = escape(line)
            escaped = _re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
            escaped = _re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"<i>\1</i>", escaped)
            escaped = _re.sub(r"`([^`]+)`", r"<code>\1</code>", escaped)
            lines.append(f"<p>{escaped}</p>")
        html = "\n".join(lines).replace("<li>", "<ul><li>").replace("</li>\n<ul>", "</li>")
        html = _re.sub(r"(<li>.*?</li>)(?!\s*<li>)", r"\1</ul>", html, flags=_re.DOTALL)
    else:
        html = "".join(f"<p>{escape(line) or '&nbsp;'}</p>" for line in body.split("\n"))

    size = max(8, min(_int(options, "font_size", 11), 24))
    family = str(options.get("font", "sans-serif"))
    css = (f"body {{ font-family: {family}; font-size: {size}px; line-height: 1.55; }}"
           f"h1,h2,h3 {{ margin: 12px 0 6px; }} code {{ font-family: monospace; }}")
    try:
        data = _story_to_pdf(html, css, str(options.get("size", "a4")),
                             max(0, min(_int(options, "margin", 56), 200)))
    except Exception as exc:  # noqa: BLE001
        return ToolResult(meta={"error": f"Could not lay that out: {exc}"})
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    return ToolResult(files=[_save_result(data, "document.pdf")],
                      meta={"pages": doc.page_count, "words": len(body.split()),
                            "bytes": len(data)})


@register("epub-to-pdf")
def epub_to_pdf(files: list[Path], text: str, options: dict) -> ToolResult:
    """Convert an EPUB to PDF.

    An EPUB is a zip of XHTML files. Those are read in the order the spine
    declares — reading them alphabetically would shuffle the chapters.
    """
    import re as _re
    import zipfile
    from html import unescape

    if not files:
        return ToolResult(meta={"error": "Upload an .epub file."})
    if not _require_fitz():
        return ToolResult(meta={"error": "PyMuPDF is not installed on the server."})
    try:
        book = zipfile.ZipFile(files[0])
    except zipfile.BadZipFile:
        return ToolResult(meta={"error": "That is not a valid EPUB — the file is not a zip."})
    names = book.namelist()
    opf = next((n for n in names if n.lower().endswith(".opf")), None)
    order: list[str] = []
    if opf:
        manifest = book.read(opf).decode("utf-8", "replace")
        base = opf.rsplit("/", 1)[0] + "/" if "/" in opf else ""
        ids = dict(_re.findall(r'<item[^>]*id="([^"]+)"[^>]*href="([^"]+)"', manifest))
        ids.update({v: k for k, v in
                    _re.findall(r'<item[^>]*href="([^"]+)"[^>]*id="([^"]+)"', manifest)})
        for ref in _re.findall(r'<itemref[^>]*idref="([^"]+)"', manifest):
            href = ids.get(ref)
            if href:
                candidate = (base + href).replace("//", "/")
                if candidate in names:
                    order.append(candidate)
    if not order:
        order = sorted(n for n in names if n.lower().endswith((".xhtml", ".html", ".htm")))
    if not order:
        return ToolResult(meta={"error": "No readable chapters found inside that EPUB."})
    chapters = []
    for name in order[:400]:
        try:
            raw = book.read(name).decode("utf-8", "replace")
        except KeyError:
            continue
        inner = _re.search(r"<body[^>]*>(.*?)</body>", raw, _re.DOTALL | _re.IGNORECASE)
        content = inner.group(1) if inner else raw
        # Images and scripts point at files inside the zip that the renderer
        # cannot reach, so they are dropped rather than left as broken boxes.
        content = _re.sub(r"<(script|style)\b.*?</\1>", "", content, flags=_re.DOTALL | _re.IGNORECASE)
        content = _re.sub(r"<img[^>]*>", "", content, flags=_re.IGNORECASE)
        chapters.append(content)
    html = "<div>" + "\n<hr>\n".join(chapters) + "</div>"
    try:
        data = _story_to_pdf(html, "body{font-family:serif;font-size:12px;line-height:1.6}",
                             str(options.get("size", "a5")),
                             max(0, min(_int(options, "margin", 40), 200)))
    except Exception as exc:  # noqa: BLE001
        return ToolResult(meta={"error": f"Could not lay that book out: {exc}"})
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    title = unescape(_stem(files[0]))
    return ToolResult(files=[_save_result(data, f"{title}.pdf")],
                      meta={"chapters": len(chapters), "pages": doc.page_count,
                            "note": "Text only — embedded images and fonts are not carried over."})


@register("pdf-to-epub")
def pdf_to_epub(files: list[Path], text: str, options: dict) -> ToolResult:
    """Build a reflowable EPUB from a PDF's text."""
    import io as _io
    import re as _re
    import zipfile
    from html import escape

    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    title = str(options.get("title", "")).strip() or _stem(files[0])
    author = str(options.get("author", "")).strip() or "Unknown"
    chapters, current = [], []
    for page in doc:
        content = page.get_text().strip()
        if content:
            current.append(content)
        # Split on the page interval rather than guessing chapters from headings,
        # which is unreliable across layouts.
        if len(current) >= max(1, _int(options, "pages_per_chapter", 10)):
            chapters.append("\n\n".join(current)); current = []
    if current:
        chapters.append("\n\n".join(current))
    if not chapters:
        return ToolResult(meta={"error": "No text found — a scanned PDF needs OCR first."})

    buf = _io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as epub:
        # mimetype must be first and STORED, uncompressed — readers check that.
        epub.writestr(zipfile.ZipInfo("mimetype"), "application/epub+zip",
                      compress_type=zipfile.ZIP_STORED)
        epub.writestr("META-INF/container.xml",
                      '<?xml version="1.0"?><container version="1.0" '
                      'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
                      '<rootfiles><rootfile full-path="OEBPS/content.opf" '
                      'media-type="application/oebps-package+xml"/></rootfiles></container>')
        items, refs, navs = [], [], []
        for index, body in enumerate(chapters, start=1):
            paragraphs = "".join(f"<p>{escape(p)}</p>"
                                 for p in _re.split(r"\n\s*\n", body) if p.strip())
            epub.writestr(f"OEBPS/ch{index}.xhtml",
                          '<?xml version="1.0" encoding="utf-8"?>'
                          '<html xmlns="http://www.w3.org/1999/xhtml"><head>'
                          f"<title>Chapter {index}</title></head><body>"
                          f"<h2>Chapter {index}</h2>{paragraphs}</body></html>")
            items.append(f'<item id="ch{index}" href="ch{index}.xhtml" '
                         f'media-type="application/xhtml+xml"/>')
            refs.append(f'<itemref idref="ch{index}"/>')
            navs.append(f'<li><a href="ch{index}.xhtml">Chapter {index}</a></li>')
        epub.writestr("OEBPS/nav.xhtml",
                      '<?xml version="1.0" encoding="utf-8"?>'
                      '<html xmlns="http://www.w3.org/1999/xhtml" '
                      'xmlns:epub="http://www.idpf.org/2007/ops"><head><title>Contents</title>'
                      '</head><body><nav epub:type="toc"><h1>Contents</h1><ol>'
                      + "".join(navs) + "</ol></nav></body></html>")
        epub.writestr("OEBPS/content.opf",
                      '<?xml version="1.0" encoding="utf-8"?>'
                      '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" '
                      'unique-identifier="bookid"><metadata '
                      'xmlns:dc="http://purl.org/dc/elements/1.1/">'
                      f"<dc:title>{escape(title)}</dc:title>"
                      f"<dc:creator>{escape(author)}</dc:creator>"
                      '<dc:language>en</dc:language>'
                      f"<dc:identifier id=\"bookid\">urn:uuid:{abs(hash(title)):032x}</dc:identifier>"
                      "</metadata><manifest>"
                      '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" '
                      'properties="nav"/>' + "".join(items) +
                      "</manifest><spine>" + "".join(refs) + "</spine></package>")
    data = buf.getvalue()
    return ToolResult(files=[_save_named(data, f"{title}.epub")],
                      meta={"chapters": len(chapters), "source_pages": doc.page_count,
                            "bytes": len(data),
                            "note": "Text only — an EPUB reflows, so the PDF's layout is not kept."})


# ===========================================================================
# Forms
# ===========================================================================

@register("pdf-extract-form-data")
def pdf_extract_form_data(files: list[Path], text: str, options: dict) -> ToolResult:
    import json as _json

    guard = _need_pdf(files)
    if guard:
        return guard
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    fields = []
    for page in doc:
        for widget in page.widgets() or []:
            fields.append({"page": page.number + 1,
                           "name": widget.field_name,
                           "type": widget.field_type_string,
                           "value": widget.field_value,
                           "options": list(widget.choice_values or []) or None,
                           "required": bool(widget.field_flags & 2)})
    if not fields:
        return ToolResult(meta={"error": "This PDF has no fillable form fields."})
    filled = sum(1 for f in fields if f["value"] not in (None, "", "Off"))
    body = _json.dumps({f["name"]: f["value"] for f in fields}, indent=2, ensure_ascii=False)
    return ToolResult(text=body,
                      files=[_save_named(body.encode(), f"formdata_{_stem(files[0])}.json")],
                      meta={"fields": len(fields), "filled": filled,
                            "empty": len(fields) - filled, "detail": fields})


@register("pdf-fill-form")
def pdf_fill_form(files: list[Path], text: str, options: dict) -> ToolResult:
    """Fill a PDF's form fields from JSON, or from name=value lines."""
    import json as _json

    guard = _need_pdf(files)
    if guard:
        return guard
    raw = (text or "").strip()
    if not raw:
        return ToolResult(meta={
            "error": 'Enter the values as JSON — {"name": "Ann"} — or one name=value per line. '
                     "Run Extract Form Data first to see the field names."
        })
    try:
        values = _json.loads(raw)
        if not isinstance(values, dict):
            raise ValueError
    except (ValueError, _json.JSONDecodeError):
        values = {}
        for line in raw.split("\n"):
            if "=" in line:
                key, _, value = line.partition("=")
                values[key.strip()] = value.strip()
        if not values:
            return ToolResult(meta={"error": "Could not read any field values from that."})
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    updated, missing = [], set(values)
    for page in doc:
        for widget in page.widgets() or []:
            if widget.field_name in values:
                new = values[widget.field_name]
                if widget.field_type_string in ("CheckBox", "RadioButton"):
                    widget.field_value = str(new).lower() in {"1", "true", "yes", "on", "x"}
                else:
                    widget.field_value = str(new)
                widget.update()
                updated.append(widget.field_name)
                missing.discard(widget.field_name)
    if not updated:
        return ToolResult(meta={"error": "None of those names match a field in this PDF."})
    if _flag(options, "flatten"):
        # Flattening makes the answers permanent — the right thing before sending
        # a completed form, and the wrong thing if it still needs editing.
        try:
            doc.bake()
        except Exception:  # noqa: BLE001
            pass
    return ToolResult(files=[_save_result(doc.tobytes(), f"filled_{_stem(files[0])}.pdf")],
                      meta={"fields_filled": len(updated), "filled": updated,
                            "not_found": sorted(missing) or "none",
                            "flattened": _flag(options, "flatten")})


@register("pdf-create-form")
def pdf_create_form(files: list[Path], text: str, options: dict) -> ToolResult:
    """Add fillable fields to a PDF, one per line as `Label: type`."""
    guard = _need_pdf(files)
    if guard:
        return guard
    import fitz

    lines = [ln.strip() for ln in (text or "").split("\n") if ln.strip()]
    if not lines:
        return ToolResult(meta={
            "error": "One field per line, as  Full name: text  — types: text, checkbox, multiline."
        })
    try:
        doc = _open_pdf(files[0])
    except ValueError as e:
        return ToolResult(meta={"error": str(e)})
    page_number = max(1, min(_int(options, "page", 1), doc.page_count))
    page = doc[page_number - 1]
    top = _int(options, "start_y", 100)
    spacing = max(28, _int(options, "spacing", 44))
    label_width = _int(options, "label_width", 150)
    field_width = _int(options, "field_width", 240)
    created = []
    for index, line in enumerate(lines):
        label, _, kind = line.partition(":")
        label, kind = label.strip(), (kind.strip().lower() or "text")
        y = top + index * spacing
        if y > page.rect.height - 60:
            break  # ran off the page rather than drawing outside it
        page.insert_text((60, y + 12), f"{label}:", fontsize=10, fontname="helv")
        widget = fitz.Widget()
        widget.field_name = label
        widget.rect = fitz.Rect(60 + label_width, y, 60 + label_width + field_width, y + 20)
        if kind == "checkbox":
            widget.field_type = fitz.PDF_WIDGET_TYPE_CHECKBOX
            widget.rect = fitz.Rect(60 + label_width, y, 60 + label_width + 16, y + 16)
        else:
            widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
            if kind == "multiline":
                widget.field_flags = 4096  # Ff bit 13 = multiline
                widget.rect = fitz.Rect(60 + label_width, y, 60 + label_width + field_width, y + 56)
        widget.border_color = (0.6, 0.6, 0.6)
        widget.border_width = 0.8
        widget.fill_color = (0.97, 0.97, 1)
        page.add_widget(widget)
        created.append({"name": label, "type": kind})
    if not created:
        return ToolResult(meta={"error": "No room on that page for any fields."})
    return ToolResult(files=[_save_result(doc.tobytes(), f"form_{_stem(files[0])}.pdf")],
                      meta={"fields_created": len(created), "fields": created,
                            "page": page_number,
                            "skipped": len(lines) - len(created) or None})
